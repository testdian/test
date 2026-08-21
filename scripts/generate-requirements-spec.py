#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从 prototypeRequirements.ts 生成需规 Word 文档（含页面截图）。
用法：先 npm run start，再 python3 scripts/generate-requirements-spec.py
"""
from __future__ import annotations

import json
import re
import subprocess
import sys
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
REQ_TS = ROOT / "src/config/prototypeRequirements.ts"
SCREENSHOT_TS = ROOT / "src/config/prototypeRequirementScreenshots.ts"
SCREENSHOT_DIR = ROOT / "public/prototype-requirements"
OUTPUT_DOCX = ROOT / "docs/碳管理平台原型需求规格说明书.docx"
BASE_URL = "http://localhost:3005"

# ---------- 解析 TS 配置（与源码保持一致，不编造条目） ----------

def requirement_screenshot_slug(menu: str, feature: str) -> str:
    slug = f"{menu}-{feature}"
    slug = re.sub(r'[/\\>|：:?*"<>]', "-", slug)
    slug = re.sub(r"\s+", "-", slug)
    slug = re.sub(r"-+", "-", slug).strip("-")
    return slug


def parse_requirements() -> list[dict]:
    text = REQ_TS.read_text(encoding="utf-8")
    items: list[dict] = []
    pattern = re.compile(
        r"\{\s*menu:\s*'((?:\\'|[^'])*)',\s*feature:\s*'((?:\\'|[^'])*)',\s*description:\s*\n?\s*'((?:\\'|[^'])*)'",
        re.MULTILINE,
    )
    for match in pattern.finditer(text):
        menu, feature, description = match.groups()
        items.append(
            {
                "menu": menu.replace("\\'", "'"),
                "feature": feature.replace("\\'", "'"),
                "description": description.replace("\\'", "'"),
            }
        )
    return items


def parse_screenshot_maps() -> tuple[dict[str, str], dict[str, str]]:
    text = SCREENSHOT_TS.read_text(encoding="utf-8")
    overrides: dict[str, str] = {}
    menu_defaults: dict[str, str] = {}

    override_block = re.search(
        r"REQUIREMENT_SCREENSHOT_OVERRIDES[^=]*=\s*\{([^}]*)\}",
        text,
        re.DOTALL,
    )
    if override_block:
        for key, path in re.findall(
            r"'([^']+)':\s*'([^']+)'", override_block.group(1)
        ):
            overrides[key] = path

    menu_block = re.search(
        r"MENU_DEFAULT_SCREENSHOTS[^=]*=\s*\{([\s\S]*?)\n\};",
        text,
    )
    if menu_block:
        for quoted, plain, path in re.findall(
            r"(?:'([^']+')|([\u4e00-\u9fffA-Za-z0-9_-]+))\s*:\s*'([^']+)'",
            menu_block.group(1),
        ):
            menu_key = (quoted or plain or "").strip("'")
            if menu_key:
                menu_defaults[menu_key] = path

    return overrides, menu_defaults


def parse_menu_doc_screenshots() -> dict[str, list[dict[str, str]]]:
    """解析 MENU_DOC_SCREENSHOTS：菜单 -> [{caption, path}, ...]"""
    text = SCREENSHOT_TS.read_text(encoding="utf-8")
    result: dict[str, list[dict[str, str]]] = {}
    block = re.search(r"MENU_DOC_SCREENSHOTS[^=]*=\s*\{([\s\S]*?)\n\};", text)
    if not block:
        return result

    menu_pattern = re.compile(r"'([^']+)':\s*\[([\s\S]*?)\]\s*,?")
    entry_pattern = re.compile(r"caption:\s*'([^']+)',\s*path:\s*'([^']+)'")
    for menu_match in menu_pattern.finditer(block.group(1)):
        menu = menu_match.group(1)
        entries = [
            {"caption": cap, "path": path}
            for cap, path in entry_pattern.findall(menu_match.group(2))
        ]
        if entries:
            result[menu] = entries
    return result


def resolve_screenshot(
    item: dict,
    overrides: dict[str, str],
    menu_defaults: dict[str, str],
) -> str:
    key = f"{item['menu']}::{item['feature']}"
    if key in overrides:
        return overrides[key]
    if item["menu"] in menu_defaults:
        return menu_defaults[item["menu"]]
    return f"/prototype-requirements/{requirement_screenshot_slug(item['menu'], item['feature'])}.png"


def resolve_menu_screenshot(
    menu: str,
    items: list[dict],
    overrides: dict[str, str],
    menu_defaults: dict[str, str],
) -> str:
    if menu in menu_defaults:
        return menu_defaults[menu]
    first = items[0]
    return resolve_screenshot(first, overrides, menu_defaults)


def resolve_menu_screenshots(
    menu: str,
    items: list[dict],
    overrides: dict[str, str],
    menu_defaults: dict[str, str],
    menu_doc_screenshots: dict[str, list[dict[str, str]]],
) -> list[dict[str, str]]:
    if menu in menu_doc_screenshots:
        return menu_doc_screenshots[menu]
    path = resolve_menu_screenshot(menu, items, overrides, menu_defaults)
    return [{"caption": "页面", "path": path}]


def group_by_menu(requirements: list[dict]) -> list[tuple[str, list[dict]]]:
    menus: list[str] = []
    grouped: dict[str, list[dict]] = {}
    for item in requirements:
        menu = item["menu"]
        if menu not in grouped:
            menus.append(menu)
            grouped[menu] = []
        grouped[menu].append(item)
    return [(menu, grouped[menu]) for menu in menus]


def format_menu_requirements(items: list[dict]) -> str:
    """将同一菜单下各功能点说明合并为一条需规（不编造，仅拼接原文）。"""
    lines: list[str] = []
    for index, item in enumerate(items, start=1):
        feature = item["feature"]
        description = item["description"]
        if len(items) == 1:
            return description
        lines.append(f"{index}. 【{feature}】{description}")
    return "\n".join(lines)


# ---------- 截图抓取 ----------

USER_KEY = "React-ant-Admin-user"
ROLE_KEY = "carbon_user_role"
MOCK_USER = {
    "accessToken": "prototype-requirements-spec",
    "username": "prototype_admin",
    "realName": "原型管理员",
    "permissions": ["*"],
    "userType": 0,
}

# 文件名 -> (路径, 角色)
CAPTURE_ROUTES: dict[str, tuple[str, str]] = {
    "登录页-Logo区域.png": ("/login", "none"),
    "全局布局-侧边栏Logo.png": ("/home", "admin"),
    "全局布局-需求说明抽屉.png": ("/home", "admin"),
    "主页-主页菜单.png": ("/home", "admin"),
    "数据看板-图表展示.png": ("/dataDashboard", "admin"),
    "系统管理-外部用户-列表与搜索.png": ("/sys/user/external", "admin"),
    "排放目标-页面整体.png": ("/emissionTarget", "admin"),
    "排放目标-菜单.png": ("/emissionTarget", "admin"),
    "排放目标-列表操作.png": ("/emissionTarget", "admin"),
    "排放目标-上年实际.png": ("/emissionTarget", "admin"),
    "排放目标-月度目标.png": ("/emissionTarget", "admin"),
    "排放目标-表格交互.png": ("/emissionTarget", "admin"),
    "减排措施-页面整体.png": ("/reductionMeasures", "admin"),
    "减排措施-菜单.png": ("/reductionMeasures", "admin"),
    "减排措施-图表.png": ("/reductionMeasures", "admin"),
    "减排措施-范围筛选.png": ("/reductionMeasures", "admin"),
    "供应链碳管理-模块说明.png": ("/sys/supplyChainCarbon/targetMgmt", "admin"),
    "供应链碳管理-菜单.png": ("/sys/supplyChainCarbon/targetMgmt", "admin"),
    "供应链碳管理-调研填报任务-列表-状态与操作.png": (
        "/sys/supplyChainCarbon/questionnaire",
        "admin",
    ),
    "供应链碳管理-调研填报任务-基础信息-所属组织.png": (
        "/sys/supplyChainCarbon/questionnaire/create",
        "admin",
    ),
    "供应链碳管理-调研填报任务-基础信息-任务名称.png": (
        "/sys/supplyChainCarbon/questionnaire/create",
        "admin",
    ),
    "供应链碳管理-调研填报任务-基础信息-任务说明.png": (
        "/sys/supplyChainCarbon/questionnaire/create",
        "admin",
    ),
    "供应链碳管理-调研填报任务-基础信息-截止日期.png": (
        "/sys/supplyChainCarbon/questionnaire/create",
        "admin",
    ),
    "供应链碳管理-调研填报任务-选择供应商.png": (
        "/sys/supplyChainCarbon/questionnaire/create",
        "admin",
    ),
    "供应链碳管理-调研填报任务-表单预览.png": (
        "/sys/supplyChainCarbon/questionnaire/create",
        "admin",
    ),
    "供应链碳管理-减排目标管理-列表搜索.png": (
        "/sys/supplyChainCarbon/targetMgmt",
        "admin",
    ),
    "供应链碳管理-减排目标管理-新增-编辑-查看表单.png": (
        "/sys/supplyChainCarbon/targetMgmt/add/0",
        "admin",
    ),
    "供应链碳管理-减排目标管理-推送至供应商.png": (
        "/sys/supplyChainCarbon/targetMgmt",
        "admin",
    ),
    "供应链碳管理-减排目标管理-目标年度.png": (
        "/sys/supplyChainCarbon/targetMgmt/add/0",
        "admin",
    ),
    "供应链碳管理-减排目标管理-减排类别.png": (
        "/sys/supplyChainCarbon/targetMgmt/add/0",
        "admin",
    ),
    "供应链碳管理-减排目标管理-产品名称.png": (
        "/sys/supplyChainCarbon/targetMgmt/add/0",
        "admin",
    ),
    "供应链碳管理-减排目标管理-列表状态与操作.png": (
        "/sys/supplyChainCarbon/targetMgmt",
        "admin",
    ),
    "供应链碳管理-计划审核-页面说明.png": ("/sys/supplyChainCarbon/plans", "admin"),
    "供应链碳管理-计划审核-列表字段.png": ("/sys/supplyChainCarbon/plans", "admin"),
    "供应链碳管理-计划审核-列表搜索.png": ("/sys/supplyChainCarbon/plans", "admin"),
    "供应链碳管理-计划审核-列表操作.png": ("/sys/supplyChainCarbon/plans", "admin"),
    "供应链碳管理-进度追踪看板-组织碳搜索.png": (
        "/sys/supplyChainCarbon/progress",
        "admin",
    ),
    "供应链碳管理-进度追踪看板-组织碳图表.png": (
        "/sys/supplyChainCarbon/progress",
        "admin",
    ),
    "供应链碳管理-进度追踪看板-组织碳表格.png": (
        "/sys/supplyChainCarbon/progress",
        "admin",
    ),
    "供应链碳管理-进度追踪看板-产品碳搜索.png": (
        "/sys/supplyChainCarbon/progress",
        "admin",
    ),
    "供应链碳管理-进度追踪看板-产品碳图表.png": (
        "/sys/supplyChainCarbon/progress",
        "admin",
    ),
    "供应链碳管理-进度追踪看板-产品碳表格.png": (
        "/sys/supplyChainCarbon/progress",
        "admin",
    ),
    "供应链碳管理-碳资质认证-模块说明.png": (
        "/sys/supplyChainCarbon/certificates",
        "admin",
    ),
    "供应商门户-角色切换.png": ("/home", "admin"),
    "供应商门户-主页-培训资料列表.png": ("/sys/supplierPortal/workbench", "supplierA"),
    "供应商门户-减排目标-目标确认.png": ("/sys/supplierPortal/targets", "supplierA"),
    "供应商门户-减排目标-列表与操作.png": ("/sys/supplierPortal/targets", "supplierA"),
    "供应商门户-减排目标-详情字段.png": ("/sys/supplierPortal/targets", "supplierA"),
    "供应商门户-减排计划-计划编制.png": ("/sys/supplierPortal/plans", "supplierA"),
    "供应商门户-减排计划-计划列表.png": ("/sys/supplierPortal/plans", "supplierA"),
    "供应商门户-进度上报-进度填报.png": ("/sys/supplierPortal/progress", "supplierA"),
    "供应商门户-调研填报任务-问卷填报.png": (
        "/sys/supplierPortal/questionnaire",
        "supplierA",
    ),
    "供应商门户-资质证书-证书管理.png": (
        "/sys/supplierPortal/certificates",
        "supplierA",
    ),
    "供应商门户-培训中心-培训浏览.png": ("/sys/supplierPortal/workbench", "supplierA"),
    "基础配置-线上培训管理-列表页.png": ("/sys/basicConfig/training", "admin"),
    "基础配置-线上培训管理-资料名称.png": ("/sys/basicConfig/training", "admin"),
    "基础配置-线上培训管理-内容摘要.png": ("/sys/basicConfig/training", "admin"),
    "基础配置-线上培训管理-培训内容.png": ("/sys/basicConfig/training", "admin"),
    "基础配置-调研表单配置-列表页.png": ("/sys/basicConfig/formTemplates", "admin"),
    "基础配置-调研表单配置-新增模板-供应商类别.png": (
        "/sys/basicConfig/formTemplates/create",
        "admin",
    ),
    "基础配置-调研表单配置-新增模板-模板名称.png": (
        "/sys/basicConfig/formTemplates/create",
        "admin",
    ),
    "基础配置-调研表单配置-配置字段-分区名称.png": (
        "/sys/basicConfig/formTemplates",
        "admin",
    ),
    "基础配置-调研表单配置-配置字段-字段类型.png": (
        "/sys/basicConfig/formTemplates",
        "admin",
    ),
    "基础配置-调研表单配置-配置字段-单位.png": (
        "/sys/basicConfig/formTemplates",
        "admin",
    ),
}


def capture_screenshots(paths: set[str], *, force: bool = False) -> None:
    """调用 Node + Playwright 抓取截图；失败时不阻断文档生成。"""
    targets = sorted(paths)
    if not force:
        targets = [p for p in targets if not (SCREENSHOT_DIR / p).exists()]
    if not targets:
        print(f"截图已齐全（{len(paths)} 张）")
        return

    print(f"待抓取截图 {len(targets)} 张，调用 capture-all-screenshots.mjs …")
    node_script = ROOT / "scripts/capture-all-screenshots.mjs"
    try:
        subprocess.run(
            ["node", str(node_script), *targets],
            cwd=str(ROOT),
            check=False,
            timeout=600,
        )
    except Exception as exc:  # noqa: BLE001
        print(f"截图抓取失败（将继续生成文档）: {exc}")

    still_missing = [p for p in targets if not (SCREENSHOT_DIR / p).exists()]
    if still_missing:
        print(f"仍缺失 {len(still_missing)} 张截图，文档中将标注缺失路径")
    else:
        print("截图抓取完成")


# ---------- 生成 Word ----------

def build_docx(
    menu_groups: list[tuple[str, list[dict]]],
    overrides,
    menu_defaults,
    menu_doc_screenshots,
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml.ns import qn

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11)

    title = doc.add_heading("碳管理平台原型需求规格说明书", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("文档版本：V1.2")
    doc.add_paragraph(f"编制日期：{date.today().isoformat()}")
    doc.add_paragraph(
        "文档说明：本文档由原型「需求说明」抽屉中的更新说明整理而成，按菜单顺序排列；"
        "每个菜单合并为一条完整需求说明，并附相关页面截图（含左侧菜单栏），供开发对照实现。"
    )
    doc.add_paragraph(
        "标注说明：原型界面中黄色位置图标（ModifyNote）表示该处有定制需求，"
        "点击可查看与本文档一致的说明文字。"
    )

    doc.add_heading("目录", level=1)
    for index, (menu, _) in enumerate(menu_groups, start=1):
        doc.add_paragraph(f"{index}. {menu}", style="List Number")

    doc.add_page_break()

    for menu_index, (menu, items) in enumerate(menu_groups, start=1):
        doc.add_heading(f"{menu_index}. {menu}", level=1)

        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        rows = [
            ("菜单", menu),
            ("需求说明", format_menu_requirements(items)),
        ]
        for row_index, (label, value) in enumerate(rows):
            table.rows[row_index].cells[0].text = label
            table.rows[row_index].cells[1].text = value
            for cell in table.rows[row_index].cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    para.paragraph_format.line_spacing = 1.25
                    for run in para.runs:
                        run.font.size = Pt(10.5)

        shots = resolve_menu_screenshots(
            menu, items, overrides, menu_defaults, menu_doc_screenshots
        )

        p = doc.add_paragraph()
        run = p.add_run("页面截图：")
        run.bold = True

        figure_index = 0
        for shot in shots:
            shot_file = Path(shot["path"]).name
            shot_path = SCREENSHOT_DIR / shot_file
            caption_label = shot["caption"]

            if shot_path.exists():
                doc.add_picture(str(shot_path), width=Cm(16))
                figure_index += 1
                cap = doc.add_paragraph(
                    f"图 {menu_index}-{figure_index}：{menu}（{caption_label}）"
                )
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if cap.runs:
                    cap.runs[0].font.size = Pt(9)
                    cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            else:
                doc.add_paragraph(
                    f"（截图文件缺失：public/prototype-requirements/{shot_file}，"
                    "请补充后重新生成文档）"
                )

        doc.add_paragraph("")

    doc.save(str(OUTPUT_DOCX))
    print(f"已生成：{OUTPUT_DOCX}（共 {len(menu_groups)} 个菜单）")


def main():
    requirements = parse_requirements()
    if not requirements:
        print("未能解析需求条目，请检查 prototypeRequirements.ts")
        sys.exit(1)

    overrides, menu_defaults = parse_screenshot_maps()
    menu_doc_screenshots = parse_menu_doc_screenshots()
    menu_groups = group_by_menu(requirements)
    print(f"共解析需求 {len(requirements)} 条，合并为 {len(menu_groups)} 个菜单")

    needed_files: set[str] = set()
    for menu, items in menu_groups:
        for shot in resolve_menu_screenshots(
            menu, items, overrides, menu_defaults, menu_doc_screenshots
        ):
            needed_files.add(Path(shot["path"]).name)

    capture_screenshots(needed_files, force=True)
    build_docx(menu_groups, overrides, menu_defaults, menu_doc_screenshots)


if __name__ == "__main__":
    main()
